"""
pip install python-docx
"""

from docx import Document  # Document代表一个Word文档
from docx.shared import Cm, Pt  # Cm 厘米，Pt磅

from docx.document import Document as Doc  # 只是另外一种导入Document的方法，和第一种等价

document = Document()

# 添加大标题
document.add_heading('快快乐乐学Python', 0)

# 添加段落
p = document.add_paragraph('Python是一门非常流行的编程语言，它')

# "run" 在文档处理中的术语含义是"连续文本片段"，表示一段具有相同格式的连续文本。
# 在 python-docx中，run是段落中的文本片段，可以独立设置样式。
run = p.add_run('简单')
run.bold = True
run.font.size = Pt(18)

p.add_run('而且')

run = p.add_run('优雅')
run.font.size = Pt(18)
run.underline = True

p.add_run("。")

# 添加一级标题
document.add_heading('H1', level=1)
# 添加带样式段落
document.add_paragraph('Intense quote', style='Intense Quote')
# 添加无序列表
document.add_paragraph('first item in unodered list', style='List Bullet')
document.add_paragraph('second item in unodered list', style='List Bullet')
# 添加有序列表
document.add_paragraph('first item in odered list', style='List Number')
document.add_paragraph('second item in unodered list', style='List Number')
# 添加图片
document.add_picture('day026/pic/python-logo.png', width=Cm(5.2))

# 添加分节符
document.add_section()

records = (
    ('骆昊', '男', '1995-5-5'),
    ('孙美丽', '女', '1992-2-2')
)

# 添加表格
table = document.add_table(rows=1, cols=3)
table.style = 'Dark List'
header_cells = table.rows[0].cells
header_cells[0].text = '姓名'
header_cells[1].text = '性别'
header_cells[2].text = '出生日期'

for name, sex, birthday in records:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = sex
    row_cells[2].text = birthday

# 分页符
document.add_page_break()

document.save('day026/demo.docx')
